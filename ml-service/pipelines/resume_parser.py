import importlib
import re
import os
import zipfile
from typing import Any

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

from pipelines.preprocessing import extract_skills_from_text, preprocess_pipeline
from pipelines.groq_service import is_groq_available, groq_json


class ResumeParseError(Exception):
    pass


def _extract_pdf_link_emails(filepath: str) -> list[str]:
    emails: list[str] = []
    if not pdfplumber:
        return emails
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                for link in page.hyperlinks or []:
                    uri = (link.get("uri") or "").strip()
                    if uri.lower().startswith("mailto:"):
                        addr = uri[7:].split("?")[0].strip()
                        normalized = _normalize_email(addr)
                        if normalized:
                            emails.append(normalized)
    except Exception:
        pass
    return emails


def extract_text_from_pdf(filepath: str) -> str:
    text = ""
    if pdfplumber:
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or page.extract_text(x_tolerance=2, y_tolerance=2)
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            pass
    if not text.strip():
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        except Exception:
            pass
    return text.strip()


def extract_text_from_docx(filepath: str) -> str:
    if not Document:
        return ""
    try:
        doc = Document(filepath)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    if ext in (".docx", ".doc"):
        return extract_text_from_docx(filepath)
    return ""


def _normalize_unicode_text(text: str) -> str:
    replacements = {
        "\uff20": "@",
        "\u00a0": " ",
        "\u200b": "",
        "\u200c": "",
        "\u200d": "",
        "\ufeff": "",
        "\uff0e": ".",
        "\u2024": ".",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text


def _extract_pdf_char_text(filepath: str, max_pages: int = 3) -> str:
    if not pdfplumber:
        return ""
    lines: list[str] = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages[:max_pages]:
                chars = page.chars or []
                if not chars:
                    continue
                grouped: dict[int, list[dict]] = {}
                for ch in chars:
                    y = round(ch.get("top", 0))
                    grouped.setdefault(y, []).append(ch)
                for y in sorted(grouped.keys()):
                    row = sorted(grouped[y], key=lambda c: c.get("x0", 0))
                    parts: list[str] = []
                    for i, ch in enumerate(row):
                        if i > 0:
                            gap = ch.get("x0", 0) - row[i - 1].get("x1", 0)
                            if gap > 1.5:
                                parts.append(" ")
                        parts.append(ch.get("text", ""))
                    line = "".join(parts).strip()
                    if line:
                        lines.append(line)
    except Exception:
        pass
    return "\n".join(lines)


def _load_pytesseract():
    try:
        return importlib.import_module("pytesseract")
    except ImportError:
        return None


def _ocr_pdf_text(filepath: str, max_pages: int = 2) -> str:
    if not pdfplumber:
        return ""
    ocr = _load_pytesseract()
    if ocr is None:
        return ""
    try:
        texts: list[str] = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages[:max_pages]:
                try:
                    pil = page.to_image(resolution=220).original
                    text = ocr.image_to_string(pil)
                    if text and text.strip():
                        texts.append(text)
                except Exception:
                    continue
        return "\n".join(texts)
    except Exception:
        return ""


def _extract_docx_link_emails(filepath: str) -> list[str]:
    emails: list[str] = []
    try:
        with zipfile.ZipFile(filepath) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        for match in re.finditer(r"mailto:([^\"'>&\s]+)", xml, flags=re.I):
            normalized = _normalize_email(match.group(1))
            if normalized:
                emails.append(normalized)
    except Exception:
        pass
    return emails


def _gather_file_email_hints(filepath: str) -> list[str]:
    ext = os.path.splitext(filepath)[1].lower()
    hints: list[str] = []
    if ext == ".pdf":
        hints.extend(_extract_pdf_link_emails(filepath))
    elif ext in (".docx", ".doc"):
        hints.extend(_extract_docx_link_emails(filepath))
    return hints


def _resolve_email(filepath: str, text_sources: list[str]) -> str | None:
    hints = _gather_file_email_hints(filepath)
    for source in text_sources:
        if not source or not source.strip():
            continue
        normalized_source = _normalize_unicode_text(source)
        found = extract_email(normalized_source, hints)
        if found:
            return found
        llm_found = _llm_extract_email(normalized_source)
        if llm_found:
            return llm_found
    return None


_INVALID_EMAIL_DOMAINS = frozenset({
    "example.com", "email.com", "domain.com", "test.com", "yoursite.com", "company.com",
})


def _normalize_email(raw: str | None) -> str | None:
    if not raw:
        return None
    cleaned = str(raw).strip().lower()
    cleaned = re.sub(r"^mailto:", "", cleaned)
    cleaned = re.sub(r"^(e-?mail|email id|email address|mail)\s*[:\-]\s*", "", cleaned, flags=re.I)
    cleaned = cleaned.strip("\"'<>()[]{} ")
    match = re.search(r"[\w.+-]+@[\w.-]+\.[a-zA-Z]{2,}", cleaned)
    if not match:
        return None
    email = re.sub(r"\s+", "", match.group(0).lower())
    domain = email.split("@", 1)[-1]
    if domain in _INVALID_EMAIL_DOMAINS:
        return None
    return email


def _deobfuscate_email_text(text: str) -> str:
    t = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", text)
    t = re.sub(r"\s*\[?\s*at\s*\]?\s*", "@", t, flags=re.I)
    t = re.sub(r"\s*\(at\)\s*", "@", t, flags=re.I)
    t = re.sub(r"\s*\[?\s*dot\s*\]?\s*", ".", t, flags=re.I)
    t = re.sub(r"\s*\(dot\)\s*", ".", t, flags=re.I)
    t = re.sub(r"\s*@\s*", "@", t)
    t = re.sub(r"(?<=[\w.+-])\s+(?=[\w.+-]*@)", "", t)
    t = re.sub(r"(?<=@)\s+(?=[\w.+-])", "", t)
    t = re.sub(r"(?<=[\w.+-])\s+(?=\.[a-zA-Z]{2,}\b)", "", t)
    return t


def _collect_email_candidates(text: str) -> list[str]:
    candidates: list[str] = []
    variants = [
        text,
        _deobfuscate_email_text(text),
        re.sub(r"\s+", " ", text),
        _deobfuscate_email_text(re.sub(r"[\n\r]+", " ", text)),
        re.sub(r"\s+", "", _deobfuscate_email_text(text)),
    ]

    patterns = [
        r"[\w.+-]+@[\w.-]+\.[a-zA-Z]{2,}",
        r"mailto:([\w.+-]+@[\w.-]+\.[a-zA-Z]{2,})",
        r"(?:e-?mail|email(?:\s+address)?|mail)\s*[:\-]?\s*([\w.+-]+@[\w.-]+\.[a-zA-Z]{2,})",
        r"[\w.+-]+\s*@\s*[\w.-]+\s*\.\s*[a-zA-Z]{2,}",
    ]

    for variant in variants:
        for pattern in patterns:
            for match in re.finditer(pattern, variant, flags=re.I):
                raw = match.group(1) if match.lastindex else match.group(0)
                normalized = _normalize_email(raw)
                if normalized:
                    candidates.append(normalized)

    lines = text.split("\n")
    for i, line in enumerate(lines):
        if re.search(r"\b(e-?mail|email(?:\s+address)?|mail id)\b", line, re.I):
            chunk = " ".join(lines[i : i + 3])
            normalized = _normalize_email(chunk) or _normalize_email(_deobfuscate_email_text(chunk))
            if normalized:
                candidates.append(normalized)

    seen: set[str] = set()
    unique: list[str] = []
    for email in candidates:
        if email not in seen:
            seen.add(email)
            unique.append(email)
    return unique


def extract_email(text: str, extra_candidates: list[str] | None = None) -> str | None:
    candidates = _collect_email_candidates(text)
    for email in extra_candidates or []:
        normalized = _normalize_email(email)
        if normalized:
            candidates.append(normalized)

    personal_domains = ("gmail.com", "yahoo.com", "outlook.com", "hotmail.com", "icloud.com", "proton.me", "live.com")
    for email in candidates:
        if any(email.endswith(f"@{d}") or email.split("@")[-1] == d for d in personal_domains):
            return email
    return candidates[0] if candidates else None


def extract_phone(text: str) -> str | None:
    patterns = [
        r"\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
        r"\b\d{10}\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    return None


def extract_name(text: str) -> str | None:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    skip_words = ("resume", "curriculum", "vitae", "cv", "profile", "summary")
    for line in lines[:8]:
        low = line.lower()
        if any(w in low for w in skip_words) and len(line.split()) <= 3:
            continue
        if "@" in line or re.search(r"\d{5,}", line):
            continue
        words = line.split()
        if 2 <= len(words) <= 5 and line[0].isupper():
            if not any(w.lower() in ("street", "avenue", "road", "http", "www") for w in words):
                return line
    return None


def extract_education(text: str) -> list[dict]:
    education = []
    edu_keywords = (
        "bachelor", "master", "phd", "b.tech", "m.tech", "b.e", "m.e", "mba",
        "bsc", "msc", "b.sc", "m.sc", "degree", "university", "college", "diploma",
    )
    lines = text.split("\n")
    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        if len(line_lower) < 4:
            continue
        if any(kw in line_lower for kw in edu_keywords):
            education.append({
                "institution": line.strip(),
                "details": lines[i + 1].strip() if i + 1 < len(lines) else "",
            })
    return education[:8]


def extract_experience(text: str) -> list[dict]:
    experience = []
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    years_match = re.search(
        r"(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s*(?:experience|exp)?",
        text.lower(),
    )
    total_years = int(years_match.group(1)) if years_match else 0

    date_range = re.compile(
        r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}\s*[-–—to]+\s*(?:present|current|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}|\d{4})",
        re.I,
    )

    for i, line in enumerate(lines):
        if date_range.search(line) or re.search(r"\d{4}\s*[-–—]\s*(?:\d{4}|present|current)", line, re.I):
            title = lines[i - 1] if i > 0 else ""
            company = line
            if title and len(title) < 80:
                experience.append({
                    "title": title,
                    "company": company,
                    "duration": company,
                    "years": total_years,
                })

    if not experience:
        section = False
        for line in lines:
            if re.search(r"\b(experience|employment|work history)\b", line, re.I):
                section = True
                continue
            if section and re.search(r"\b(education|skills|projects|certifications)\b", line, re.I):
                break
            if section and len(line) > 8 and line[0].isupper():
                experience.append({
                    "title": line[:120],
                    "company": "",
                    "years": total_years,
                })
                if len(experience) >= 6:
                    break

    if not experience and total_years:
        experience.append({"title": "Professional Experience", "years": total_years})

    return experience[:8]


def extract_certifications(text: str) -> list[str]:
    cert_keywords = ("certified", "certification", "certificate", "aws certified", "pmp", "cissp", "comptia")
    certs = []
    for line in text.split("\n"):
        if any(kw in line.lower() for kw in cert_keywords) and len(line.strip()) > 5:
            certs.append(line.strip()[:200])
    return certs[:10]


def extract_projects(text: str) -> list[dict]:
    projects = []
    for line in text.split("\n"):
        low = line.lower()
        if any(kw in low for kw in ("project", "developed", "built", "implemented", "designed")) and len(line) > 25:
            projects.append({"title": line.strip()[:100], "description": line.strip()[:300]})
    return projects[:8]


def _llm_extract_email(raw_text: str) -> str | None:
    if not is_groq_available() or len(raw_text) < 40:
        return None
    result = groq_json(
        "Extract candidate email from resume. JSON with key email (string or null). Never invent.",
        f"Find the email in this resume. Return null if none.\n\n{raw_text[:4000]}",
    )
    if isinstance(result, dict):
        return _normalize_email(result.get("email"))
    return None


def _llm_parse_resume(raw_text: str) -> dict | None:
    if not is_groq_available() or len(raw_text) < 80:
        return None
    result = groq_json(
        "Expert resume parser. Extract ONLY facts in the resume. JSON object only. Never invent data.",
        (
            "Parse resume. Return JSON: name, email, phone, skills[], education[], experience[], "
            f"certifications[], projects[], summary.\n\n{raw_text[:5000]}"
        ),
    )
    return result if isinstance(result, dict) else None


def _merge_parsed(rule: dict, llm: dict | None) -> dict:
    if not llm:
        return rule

    out = {**rule}
    for key in ("name", "phone", "summary"):
        val = llm.get(key)
        if val and str(val).strip() and str(val).lower() not in ("null", "none", "n/a"):
            out[key] = str(val).strip()

    llm_email = _normalize_email(llm.get("email"))
    if llm_email:
        out["email"] = llm_email

    if llm.get("skills") and isinstance(llm["skills"], list):
        merged = list(dict.fromkeys([*out.get("skills", []), *[str(s) for s in llm["skills"] if s]]))
        out["skills"] = merged[:30]

    for field in ("education", "experience", "certifications", "projects"):
        if llm.get(field) and isinstance(llm[field], list) and len(llm[field]) > 0:
            out[field] = llm[field]

    return out


def parse_resume(filepath: str) -> dict[str, Any]:
    if not os.path.isfile(filepath):
        raise ResumeParseError(f"Resume file not found: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()
    raw_text = _normalize_unicode_text(extract_text(filepath))
    char_text = _normalize_unicode_text(_extract_pdf_char_text(filepath)) if ext == ".pdf" else ""
    ocr_text = _normalize_unicode_text(_ocr_pdf_text(filepath)) if ext == ".pdf" else ""

    combined_text = "\n".join(part for part in (raw_text, char_text, ocr_text) if part).strip()
    if not combined_text or len(combined_text) < 50:
        raise ResumeParseError(
            "Could not extract readable text from this resume. "
            "Use a text-based PDF or DOCX (not a scanned image-only PDF)."
        )

    primary_text = raw_text or char_text or ocr_text or combined_text
    processed = preprocess_pipeline(primary_text)
    llm_data = _llm_parse_resume(combined_text)

    resolved_email = _resolve_email(filepath, [primary_text, char_text, ocr_text, combined_text])

    rule_based = {
        "name": extract_name(primary_text) or extract_name(combined_text),
        "email": resolved_email,
        "phone": extract_phone(combined_text) or extract_phone(primary_text),
        "education": extract_education(combined_text),
        "experience": extract_experience(combined_text),
        "skills": extract_skills_from_text(combined_text),
        "certifications": extract_certifications(combined_text),
        "projects": extract_projects(combined_text),
        "raw_text": combined_text[:8000],
        "processed_text": processed[:4000],
        "summary": "",
    }

    merged = _merge_parsed(rule_based, llm_data)

    if not merged.get("name"):
        merged["name"] = extract_name(combined_text) or "Unknown (see resume)"

    if not merged.get("skills"):
        merged["skills"] = extract_skills_from_text(combined_text)

    if not merged.get("email"):
        merged["email"] = _resolve_email(filepath, [combined_text, char_text, ocr_text, primary_text])

    merged["parse_source"] = "groq+rules" if llm_data else "rules"
    if ocr_text and merged.get("email") and not extract_email(primary_text):
        merged["parse_source"] += "+ocr"
    merged["text_length"] = len(combined_text)
    return merged
